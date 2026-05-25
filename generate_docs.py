import sys
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

def create_resume_docx():
    doc = Document()
    
    # Header
    title = doc.add_heading('Brandon Dando', 0)
    subtitle = doc.add_paragraph('Full-Stack Developer | RF Engineer | Technical Mentor')
    subtitle.alignment = 1
    
    contact = doc.add_paragraph()
    contact.add_run('Nairobi County | 0112607179 | dandobrandon0@gmail.com\n')
    contact.add_run('LinkedIn: https://www.linkedin.com/in/timon-dando-b713a5261/ | GitHub: https://github.com/djBrandy')
    contact.alignment = 1

    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        "Innovative Engineer and Technical Founder operating at the intersection of high-level software architecture "
        "and low-level hardware exploration. Co-founder of ASENASS Developers, specialized in building scalable "
        "distributed systems and IoT solutions. Expert in Python/Flask and React, with a deep specialization in "
        "RF signal analysis and SDR-based security research."
    )

    # Skills
    doc.add_heading('Technical Skills', level=1)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Software Engineering:').bold = True
    p.add_run(' Python, Flask, React, SQL Alchemy, REST APIs, JavaScript (ES6+), WebSocket, Docker, CI/CD (GitHub Actions), Redis, PostgreSQL.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Hardware & RF:').bold = True
    p.add_run(' RF Engineering, SDR (Software Defined Radio), BLE Protocol Analysis, Embedded C++, Signal Processing, ESP32/Arduino, MQTT.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Leadership:').bold = True
    p.add_run(' Strategic Technical Planning, Community Building, Technical Mentoring, Agile Project Management.')

    # Experience
    doc.add_heading('Professional Experience', level=1)
    
    doc.add_heading('ASENASS Developers | Co-Founder', level=2)
    doc.add_paragraph('January 2024 – Present')
    doc.add_paragraph('Co-engineered 5+ commercial full-stack applications from concept to production using Flask and React, achieving 99.9% uptime by containerizing microservices with Docker.', style='List Bullet')
    doc.add_paragraph('Architected a custom Learning Management System (LMS) serving 500+ active students, featuring real-time interactive coding challenges and WebSocket-based progress tracking.', style='List Bullet')
    doc.add_paragraph('Optimized backend performance by implementing Redis caching and database indexing, reducing API response times by 35% for high-traffic endpoints.', style='List Bullet')
    doc.add_paragraph('Spearheaded deployment pipelines using GitHub Actions, reducing release cycles by 40% through automated testing and CI/CD integration.', style='List Bullet')

    doc.add_heading('ASENASS Academy | Lead Technical Mentor', level=2)
    doc.add_paragraph('January 2024 – Present')
    doc.add_paragraph('Designed and delivered a comprehensive Software Engineering curriculum focusing on modern backend architecture, resulting in a 40% increase in student engagement.', style='List Bullet')
    doc.add_paragraph('Mentored 10+ junior developers through 1-on-1 sessions and rigorous code reviews, focusing on clean code principles and industrial-grade toolsets (Git, Docker, CI/CD).', style='List Bullet')

    # Key Projects
    doc.add_heading('Key Projects', level=1)
    
    doc.add_heading('RF-SDR Spectral Analysis Dashboard', level=2)
    doc.add_paragraph('Built a real-time frequency monitoring tool using Flask and React. Integrated pyrtlsdr for raw I/Q sample capture and used D3.js to render low-latency waterfall visualizations and PSD graphs.', style='List Bullet')

    doc.add_heading('BLE Security Research & Signal Jammer', level=2)
    doc.add_paragraph('Developed a hardware-based interference system for BLE protocol vulnerability research. Engineered firmware in C++ for ESP32 to analyze and mitigate signal interference patterns in local IoT ecosystems.', style='List Bullet')

    # Education
    doc.add_heading('Education & Certifications', level=1)
    doc.add_paragraph('Moringa School | Professional Certificate in Software Engineering (2024)', style='List Bullet')
    doc.add_paragraph('Focus: Full-Stack Development, Agile Methodologies, and Technical Leadership.', style='List Bullet')

    doc.save('Brandon_Dando_Resume.docx')

def create_resume_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Brandon Dando", align="C")
    pdf.ln(8)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, "Nairobi County | 0112607179 | dandobrandon0@gmail.com", align="C")
    pdf.ln(5)
    pdf.cell(0, 5, "LinkedIn: linkedin.com/in/timon-dando-b713a5261/ | GitHub: github.com/djBrandy", align="C")
    pdf.ln(10)

    # Sections
    sections = [
        ("Professional Summary", "Innovative Engineer and Technical Founder operating at the intersection of high-level software architecture and low-level hardware exploration. Co-founder of ASENASS Developers, specialized in building scalable distributed systems and IoT solutions. Expert in Python/Flask and React, with a deep specialization in RF signal analysis and SDR-based security research."),
        ("Technical Skills", "- Software: Python, Flask, React, SQL Alchemy, REST APIs, Docker, CI/CD, Redis, PostgreSQL.\n- Hardware/RF: SDR, BLE Analysis, Embedded C++, Signal Processing, ESP32, MQTT.\n- Leadership: Technical Planning, Community Building, Mentoring, Agile Management."),
        ("Experience", "ASENASS Developers | Co-Founder (2024 - Present)\n- Co-engineered 5+ commercial applications using Flask/React; achieved 99.9% uptime via Docker.\n- Architected an LMS for 500+ students with real-time WebSocket-based coding challenges.\n- Reduced API response times by 35% through Redis caching and DB optimization.\n- Automated CI/CD pipelines via GitHub Actions, reducing release cycles by 40%.\n\nASENASS Academy | Lead Technical Mentor (2024 - Present)\n- Designed SE curriculum resulting in 40% increase in student engagement.\n- Mentored 10+ developers on clean code and professional toolsets."),
        ("Key Projects", "RF-SDR Dashboard: Real-time frequency monitor using Flask/React and pyrtlsdr; rendered low-latency waterfall plots via D3.js.\nBLE Security Research: Developed ESP32 firmware in C++ for signal interference and protocol analysis."),
        ("Education", "- Moringa School: Professional Certificate in Software Engineering (2024)")
    ]

    for title, content in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, title)
        pdf.ln(8)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, content)
        pdf.ln(4)

    pdf.output("Brandon_Dando_Resume.pdf")

def create_cv_docx():
    doc = Document()
    
    # Header
    doc.add_heading('Curriculum Vitae - Brandon Dando', 0)
    p = doc.add_paragraph()
    p.add_run('Nairobi County | 0112607179 | dandobrandon0@gmail.com\n').bold = True
    p.add_run('LinkedIn: https://www.linkedin.com/in/timon-dando-b713a5261/\n')
    p.add_run('GitHub: https://github.com/djBrandy\n')
    p.add_run('Certificate: https://moringa.my.salesforce-sites.com/certificateStatus?id=a0PQ200000ADeYW')
    p.alignment = 1

    # Profile
    doc.add_heading('Professional Profile', level=1)
    doc.add_paragraph(
        "A highly technical and versatile engineer with a unique blend of expertise in high-level software development "
        "and low-level hardware experimentation. As a Co-Founder and Lead Technical Mentor, I have demonstrated "
        "excellence in building both complex technical ecosystems and high-performing engineering talent. "
        "My current work focuses on bridging the gap between digital software elegance and physical RF signal analysis."
    )

    # Areas of Expertise
    doc.add_heading('Areas of Expertise', level=1)
    doc.add_paragraph('Distributed Systems: Expert-level Python/Flask backend architecture and containerized microservices.', style='List Bullet')
    doc.add_paragraph('RF & Wireless Security: Designing and hacking wireless systems using SDR, BLE, and sub-GHz protocols.', style='List Bullet')
    doc.add_paragraph('API Engineering: Crafting scalable, secure, and well-documented RESTful frameworks with JWT and Redis.', style='List Bullet')
    doc.add_paragraph('Technical Leadership: Founding startups, curriculum design, and managing cross-functional technical teams.', style='List Bullet')

    # Experience
    doc.add_heading('Professional Experience', level=1)
    
    doc.add_heading('ASENASS Developers | Co-Founder', level=2)
    doc.add_paragraph('January 2024 – Present')
    doc.add_paragraph('Commercial Delivery: Co-engineered 5+ commercial web applications from concept to production, maintaining a 99.9% uptime record.', style='List Bullet')
    doc.add_paragraph('Infrastructure: Architected the core distributed systems for the ASENASS ecosystem using Docker and GitHub Actions for CI/CD.', style='List Bullet')
    doc.add_paragraph('Performance: Optimized API performance by 35% through advanced caching strategies and database normalization.', style='List Bullet')

    doc.add_heading('ASENASS Academy | Lead Technical Mentor', level=2)
    doc.add_paragraph('January 2024 – Present')
    doc.add_paragraph('Curriculum Engineering: Developed comprehensive training modules for Software Engineering, increasing student engagement by 40%.', style='List Bullet')
    doc.add_paragraph('Talent Development: Mentored over 50 aspiring engineers, focusing on industrial-grade backend engineering and devops.', style='List Bullet')

    # Technical Projects (In-Depth)
    doc.add_heading('Technical Projects (In-Depth)', level=1)
    
    doc.add_heading('ASENASS Learning Management System (LMS)', level=2)
    doc.add_paragraph('Engineered a high-concurrency platform using Python/Flask and React. Implemented WebSockets for real-time interactivity, supporting 500+ concurrent learners.', style='List Bullet')

    doc.add_heading('RF-SDR Spectral Analysis Framework', level=2)
    doc.add_paragraph('Built an end-to-end framework for RF signal monitoring. Used Flask for FFT processing of raw I/Q samples and D3.js for real-time waterfall visualization.', style='List Bullet')

    doc.add_heading('BLE Security Research System', level=2)
    doc.add_paragraph('Designed a system for Bluetooth Low Energy protocol analysis using ESP32 and C++. Identified and documented vulnerabilities in local IoT devices.', style='List Bullet')

    # Education
    doc.add_heading('Education & Certifications', level=1)
    doc.add_paragraph('Moringa School | Professional Certificate in Software Engineering (2024)', style='List Bullet')

    doc.save('Brandon_Dando_CV.docx')

def create_cv_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Curriculum Vitae - Brandon Dando", align="C")
    pdf.ln(10)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 5, "Nairobi County | 0112607179 | dandobrandon0@gmail.com", align="C")
    pdf.ln(5)
    pdf.cell(0, 5, "LinkedIn: linkedin.com/in/timon-dando-b713a5261/ | GitHub: github.com/djBrandy", align="C")
    pdf.ln(10)

    sections = [
        ("Professional Profile", "A highly technical and versatile engineer with a unique blend of expertise in high-level software development and low-level hardware experimentation. As a Co-Founder and Lead Technical Mentor, I have demonstrated excellence in building both complex technical ecosystems and high-performing engineering talent."),
        ("Areas of Expertise", "- Distributed Systems: Python/Flask and containerized microservices.\n- RF & Wireless Security: SDR, BLE, and sub-GHz protocols.\n- API Engineering: Scalable RESTful frameworks with JWT and Redis.\n- Technical Leadership: Founding startups and curriculum design."),
        ("Professional Experience", "ASENASS Developers | Co-Founder (2024 - Present)\n- Co-engineered 5+ commercial applications with 99.9% uptime.\n- Architected core infrastructure using Docker and GitHub Actions.\n- Optimized API performance by 35%.\n\nASENASS Academy | Lead Technical Mentor (2024 - Present)\n- Developed SE curriculum increasing engagement by 40%.\n- Mentored 50+ engineers in industrial backend and devops."),
        ("Technical Projects (In-Depth)", "ASENASS LMS: Flask/React/WebSockets platform for 500+ concurrent learners.\nRF-SDR Framework: End-to-end signal monitoring and FFT processing.\nBLE Security System: ESP32/C++ system for protocol vulnerability analysis."),
        ("Education", "Moringa School: Professional Certificate in Software Engineering (2024)")
    ]

    for title, content in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 10, title)
        pdf.ln(8)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, content)
        pdf.ln(5)

    pdf.output("Brandon_Dando_CV.pdf")

if __name__ == "__main__":
    create_resume_docx()
    create_resume_pdf()
    create_cv_docx()
    create_cv_pdf()
    print("Resume and CV (.docx and .pdf) generated successfully!")
